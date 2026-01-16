#!/usr/bin/env python3
"""
VICORE - Générateur de Rapport Word
Génère un rapport complet avec documentation fonctionnelle, technique, audit et recommandations.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_with_number(doc, text, level, number=None):
    """Add a numbered heading"""
    if number:
        heading = doc.add_heading(f"{number} {text}", level)
    else:
        heading = doc.add_heading(text, level)
    return heading

def create_table_with_header(doc, headers, data, col_widths=None):
    """Create a formatted table with header row"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '2E75B6')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)

    return table

def generate_report():
    doc = Document()

    # =========================================================================
    # PAGE DE GARDE
    # =========================================================================
    title = doc.add_heading('VICORE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Virtual Core Inspection System')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].font.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    main_title = doc.add_paragraph('Rapport de Rétro-Documentation')
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    main_title.runs[0].font.size = Pt(24)
    main_title.runs[0].font.bold = True

    doc.add_paragraph()

    sub_sections = doc.add_paragraph('Documentation Fonctionnelle • Documentation Technique\nAudit de Code • Recommandations')
    sub_sections.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_sections.runs[0].font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Info box
    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ('Version', '1.0.0.12'),
        ('Date', datetime.now().strftime('%d/%m/%Y')),
        ('Projet', 'Eurotunnel Web - Inspection des Ressorts'),
        ('Auteur', 'Claude Opus 4.5 (Anthropic)')
    ]
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        info_table.rows[i].cells[1].text = value

    doc.add_page_break()

    # =========================================================================
    # TABLE DES MATIÈRES
    # =========================================================================
    doc.add_heading('Table des Matières', 1)

    toc_items = [
        ('1', 'Documentation Fonctionnelle', '3'),
        ('1.1', 'Vue d\'Ensemble du Système', '3'),
        ('1.2', 'Cas d\'Utilisation', '4'),
        ('1.3', 'Workflow Utilisateur', '6'),
        ('1.4', 'Règles Métier', '7'),
        ('2', 'Documentation Technique', '8'),
        ('2.1', 'Architecture Applicative', '8'),
        ('2.2', 'Stack Technologique', '9'),
        ('2.3', 'Structure du Code', '10'),
        ('2.4', 'API Reference', '12'),
        ('2.5', 'Modèle de Données', '14'),
        ('3', 'Audit de Code', '16'),
        ('3.1', 'Analyse Statique', '16'),
        ('3.2', 'Audit de Sécurité', '17'),
        ('3.3', 'Qualité du Code', '19'),
        ('4', 'Recommandations', '21'),
        ('4.1', 'Matrice Impact vs Effort', '21'),
        ('4.2', 'Plan d\'Action Priorisé', '23'),
    ]

    toc_table = doc.add_table(rows=len(toc_items), cols=3)
    for i, (num, title, page) in enumerate(toc_items):
        toc_table.rows[i].cells[0].text = num
        toc_table.rows[i].cells[1].text = title
        toc_table.rows[i].cells[2].text = page
        if '.' not in num:
            for cell in toc_table.rows[i].cells:
                cell.paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # =========================================================================
    # 1. DOCUMENTATION FONCTIONNELLE
    # =========================================================================
    doc.add_heading('1. Documentation Fonctionnelle', 1)

    # 1.1 Vue d'ensemble
    doc.add_heading('1.1 Vue d\'Ensemble du Système', 2)

    doc.add_heading('Description', 3)
    doc.add_paragraph(
        'VICORE (Virtual Core Inspection System) est une application web dédiée à la visualisation '
        'et à la validation des inspections automatiques de ressorts sur les trains ferroviaires. '
        'Le système est déployé pour Eurotunnel et permet aux opérateurs de surveiller en temps réel '
        'l\'état des ressorts de suspension des wagons.'
    )

    doc.add_heading('Objectifs Métier', 3)
    objectives_data = [
        ('Visualisation', 'Afficher les résultats d\'inspection des trains en temps réel'),
        ('Analyse', 'Permettre l\'examen détaillé des images de ressorts par wagon'),
        ('Confirmation', 'Valider ou infirmer les détections automatiques de ressorts manquants'),
        ('Monitoring', 'Surveiller la santé et la disponibilité des systèmes d\'inspection'),
        ('Traçabilité', 'Enregistrer les confirmations humaines pour audit')
    ]
    create_table_with_header(doc, ['Objectif', 'Description'], objectives_data)

    doc.add_paragraph()
    doc.add_heading('Périmètre Fonctionnel', 3)
    scope_items = [
        'Authentification des opérateurs',
        'Visualisation des passages de trains par installation',
        'Navigation dans les wagons et visualisation des ressorts',
        'Confirmation manuelle des ressorts détectés comme manquants',
        'Monitoring de la santé des systèmes d\'acquisition',
        'Ajustement visuel des images (contraste, luminosité)'
    ]
    for item in scope_items:
        doc.add_paragraph(item, style='List Bullet')

    # 1.2 Cas d'utilisation
    doc.add_heading('1.2 Cas d\'Utilisation', 2)

    # UC-01
    doc.add_heading('UC-01: Authentification Utilisateur', 3)
    doc.add_paragraph('Acteur: Opérateur').runs[0].bold = True
    doc.add_paragraph('Précondition: L\'utilisateur possède un compte valide dans le système.')
    doc.add_paragraph('Postcondition: L\'utilisateur a accès à l\'interface principale.')
    doc.add_paragraph('Flux Principal:')
    uc01_steps = [
        'L\'utilisateur accède à l\'URL de l\'application',
        'Le système affiche la page de connexion',
        'L\'utilisateur saisit son nom d\'utilisateur et mot de passe',
        'Le système valide les identifiants via bcrypt',
        'Le système crée une session et redirige vers la vue système'
    ]
    for i, step in enumerate(uc01_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    doc.add_paragraph('Flux Alternatif: Si les identifiants sont invalides, afficher un message d\'erreur.')

    # UC-02
    doc.add_heading('UC-02: Visualisation des Passages de Trains', 3)
    doc.add_paragraph('Acteur: Opérateur').runs[0].bold = True
    doc.add_paragraph('Précondition: Utilisateur authentifié.')
    doc.add_paragraph('Flux Principal:')
    uc02_steps = [
        'Le système charge et affiche les 50 derniers passages',
        'Les passages sont groupés par installation (Coquelles, Folkestone, etc.)',
        'Chaque passage affiche: date/heure, identifiant RFID, code de confiance',
        'L\'utilisateur sélectionne un passage',
        'Le système affiche les wagons du train avec leurs indicateurs de couleur',
        'L\'utilisateur peut charger plus de passages via le lazy loading'
    ]
    for i, step in enumerate(uc02_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    # UC-03
    doc.add_heading('UC-03: Inspection d\'un Wagon', 3)
    doc.add_paragraph('Acteur: Opérateur').runs[0].bold = True
    doc.add_paragraph('Précondition: Passage de train sélectionné.')
    doc.add_paragraph('Flux Principal:')
    uc03_steps = [
        'L\'utilisateur clique sur l\'icône d\'un wagon',
        'Le système affiche la vue détaillée avec le schéma SVG du bogie',
        'L\'image du premier ressort s\'affiche',
        'L\'utilisateur navigue entre les ressorts avec les boutons Previous/Next',
        'L\'utilisateur peut ajuster le contraste et la luminosité de l\'image',
        'Les ressorts problématiques sont signalés par des indicateurs colorés'
    ]
    for i, step in enumerate(uc03_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    # UC-04
    doc.add_heading('UC-04: Confirmation de Ressort Manquant', 3)
    doc.add_paragraph('Acteur: Opérateur').runs[0].bold = True
    doc.add_paragraph('Précondition: Ressort avec niveau de confiance faible (code > 2).')
    doc.add_paragraph('Flux Principal:')
    uc04_steps = [
        'Le système affiche le bouton "Confirm spring missing"',
        'L\'utilisateur clique sur le bouton',
        'Une fenêtre modale affiche toutes les images disponibles du ressort',
        'L\'utilisateur examine les images et choisit:',
        '   a) "Spring Present" - Le ressort est visible',
        '   b) "Spring Missing" - Le ressort est effectivement absent',
        '   c) "Unsure - Close" - Fermer sans confirmation',
        'Le système enregistre la décision et met à jour l\'affichage'
    ]
    for i, step in enumerate(uc04_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_page_break()

    # 1.3 Workflow
    doc.add_heading('1.3 Workflow Utilisateur', 2)

    doc.add_paragraph(
        'Le workflow principal de l\'application suit un parcours linéaire depuis '
        'l\'authentification jusqu\'à la confirmation des ressorts problématiques.'
    )

    workflow_data = [
        ('1', 'Login', 'Authentification de l\'opérateur', '/login'),
        ('2', 'System View', 'Liste des passages par installation', '/listpasses'),
        ('3', 'Sélection', 'Choix d\'un passage de train', 'Click sur passage'),
        ('4', 'Car View', 'Vue détaillée d\'un wagon', '/cars/{id}/{num}'),
        ('5', 'Navigation', 'Parcours des ressorts', 'Boutons Prev/Next'),
        ('6', 'Confirmation', 'Validation des ressorts AWOL', 'Modal de confirmation'),
        ('7', 'Retour', 'Retour à la liste', 'Bouton Back')
    ]
    create_table_with_header(doc, ['Étape', 'Écran', 'Action', 'Route/Trigger'], workflow_data)

    # 1.4 Règles Métier
    doc.add_heading('1.4 Règles Métier', 2)

    doc.add_heading('Niveaux de Confiance', 3)
    doc.add_paragraph(
        'Le système utilise des codes de confiance pour indiquer la fiabilité de la détection automatique:'
    )

    confidence_data = [
        ('0', 'Confirmé Présent', 'Bleu', 'Validation humaine: ressort présent'),
        ('1', 'GREEN', 'Vert', 'Haute confiance - Aucune action requise'),
        ('2', 'AMBER', 'Orange', 'Confiance moyenne - Vérification recommandée'),
        ('3', 'RED', 'Rouge', 'Faible confiance - Confirmation requise'),
        ('5', 'Non vérifié', 'Gris', 'Wagon non inspecté')
    ]
    create_table_with_header(doc, ['Code', 'Niveau', 'Couleur', 'Description'], confidence_data)

    doc.add_paragraph()
    doc.add_heading('Règles de Confirmation', 3)
    rules = [
        'Une confirmation ne peut être effectuée qu\'une seule fois (HTTP 409 si déjà confirmé)',
        'Seuls les utilisateurs authentifiés peuvent confirmer',
        'La confirmation enregistre: spring_id, statut (présent/absent), utilisateur, timestamp',
        'Un ressort confirmé comme présent passe au code de confiance 0 (bleu)'
    ]
    for rule in rules:
        doc.add_paragraph(rule, style='List Bullet')

    doc.add_page_break()

    # =========================================================================
    # 2. DOCUMENTATION TECHNIQUE
    # =========================================================================
    doc.add_heading('2. Documentation Technique', 1)

    # 2.1 Architecture
    doc.add_heading('2.1 Architecture Applicative', 2)

    doc.add_paragraph(
        'L\'application suit une architecture web classique en trois couches avec '
        'une couche de cache Redis pour optimiser les performances.'
    )

    arch_data = [
        ('Présentation', 'Vue.js 2.6, Element-UI, Bootstrap 5', 'Templates Jinja2, JavaScript'),
        ('Application', 'Flask 3.x, Python 3.10+', 'Endpoints REST, Sessions'),
        ('Données', 'PostgreSQL, SQLAlchemy 2.0', 'ORM, Fonctions stockées'),
        ('Cache', 'Redis', 'Heartbeats, Display names')
    ]
    create_table_with_header(doc, ['Couche', 'Technologies', 'Responsabilités'], arch_data)

    doc.add_paragraph()
    doc.add_heading('Diagramme d\'Architecture', 3)
    doc.add_paragraph(
        '[Browser/Vue.js] → [Flask Backend] → [PostgreSQL]\n'
        '                         ↓\n'
        '                     [Redis Cache]\n'
        '                         ↓\n'
        '                  [File System/Images]'
    )

    # 2.2 Stack
    doc.add_heading('2.2 Stack Technologique', 2)

    doc.add_heading('Backend', 3)
    backend_stack = [
        ('Python', '>=3.10, <3.13', 'Langage principal'),
        ('Flask', '>3.0', 'Framework web'),
        ('SQLAlchemy', '2.0.29', 'ORM base de données'),
        ('Pydantic', '2.9.2', 'Validation des données'),
        ('Redis', '5.1.1', 'Cache'),
        ('Bcrypt', '4.1.2', 'Hashage mots de passe'),
        ('Gunicorn', '23.0.0', 'Serveur WSGI'),
        ('Loguru', 'latest', 'Logging')
    ]
    create_table_with_header(doc, ['Package', 'Version', 'Usage'], backend_stack)

    doc.add_paragraph()
    doc.add_heading('Frontend', 3)
    frontend_stack = [
        ('Vue.js', '2.6.14', 'Framework JavaScript'),
        ('Axios', '0.21.4', 'Client HTTP'),
        ('Element-UI', '2.15.14', 'Composants UI'),
        ('Bootstrap', '5.3.3', 'Framework CSS'),
        ('jQuery', '3.7.1', 'Utilitaires DOM')
    ]
    create_table_with_header(doc, ['Librairie', 'Version', 'Usage'], frontend_stack)

    doc.add_paragraph()
    doc.add_heading('Infrastructure', 3)
    infra_stack = [
        ('PostgreSQL', '14+', 'Base de données principale'),
        ('Redis', '6+', 'Cache en mémoire'),
        ('Docker', 'latest', 'Conteneurisation'),
        ('Nginx/Traefik', '-', 'Reverse proxy (production)')
    ]
    create_table_with_header(doc, ['Composant', 'Version', 'Rôle'], infra_stack)

    doc.add_page_break()

    # 2.3 Structure du Code
    doc.add_heading('2.3 Structure du Code', 2)

    doc.add_heading('Arborescence', 3)
    structure = """
eurotunnel_web/
├── app.py                      # Point d'entrée Flask, routes principales
├── db_iface.py                 # Interface base de données, requêtes
├── user_management.py          # Authentification, gestion utilisateurs
├── system_endpoints.py         # Endpoints monitoring système
├── train_pass_endpoints.py     # Endpoints passages de trains
├── missing_spring_endpoints.py # Endpoints confirmation ressorts
├── confidence_levels.py        # Calcul niveaux de confiance
├── display_name_iface.py       # Génération noms d'affichage
├── redis_web.py                # Interface cache Redis
├── models/
│   ├── user.py                 # Modèle Pydantic utilisateur
│   └── display_names_model.py  # Modèle noms d'affichage
├── templates/                  # Templates Jinja2 (10 fichiers)
├── static/                     # Assets (CSS, JS, images)
└── tests/                      # Tests d'intégration
"""
    para = doc.add_paragraph()
    run = para.add_run(structure)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    doc.add_heading('Modules Principaux', 3)
    modules_data = [
        ('app.py', 'Configuration Flask, routes, middleware', '166 lignes'),
        ('db_iface.py', 'Requêtes SQL, conversion données', '268 lignes'),
        ('train_pass_endpoints.py', 'API passages trains, lazy loading', '117 lignes'),
        ('missing_spring_endpoints.py', 'Confirmation ressorts manquants', '55 lignes'),
        ('system_endpoints.py', 'Heartbeat, statut systèmes', '85 lignes'),
        ('confidence_levels.py', 'Calcul codes confiance', '61 lignes'),
        ('redis_web.py', 'Cache Redis', '74 lignes'),
        ('user_management.py', 'Auth bcrypt', '53 lignes')
    ]
    create_table_with_header(doc, ['Module', 'Responsabilité', 'Taille'], modules_data)

    doc.add_page_break()

    # 2.4 API Reference
    doc.add_heading('2.4 API Reference', 2)

    doc.add_heading('Endpoints Publics', 3)
    public_api = [
        ('GET', '/login', 'Affiche la page de connexion'),
        ('POST', '/login', 'Authentifie l\'utilisateur')
    ]
    create_table_with_header(doc, ['Méthode', 'Route', 'Description'], public_api)

    doc.add_paragraph()
    doc.add_heading('Endpoints Protégés (Auth requise)', 3)
    protected_api = [
        ('GET', '/', 'Redirige vers /listpasses'),
        ('GET', '/listpasses', 'Vue système principale'),
        ('GET', '/get_train_passes', 'Liste des 50 derniers passages'),
        ('POST', '/getTrainPass/{tpid}', 'Détails d\'un passage avec wagons'),
        ('POST', '/loadTrainPasses/{last}/{n}/{install?}', 'Lazy loading passages'),
        ('GET', '/cars/{tpid}/{car}', 'Vue détaillée wagon'),
        ('POST', '/getCar/{tpid}/{car}', 'Données wagon avec ressorts'),
        ('GET', '/get_all_images_for_spring/{id}', 'Images d\'un ressort'),
        ('PUT', '/put_confirmation_status/{id}/{status}', 'Confirmer ressort'),
        ('GET', '/heartbeat/{system}', 'Heartbeat système'),
        ('GET', '/hb_time/{system}', 'Timestamp heartbeat'),
        ('GET', '/get_worst_system_status', 'Statut global'),
        ('GET', '/wheelimg/{path}', 'Serveur images'),
        ('GET', '/logout', 'Déconnexion')
    ]
    create_table_with_header(doc, ['Méthode', 'Route', 'Description'], protected_api)

    doc.add_paragraph()
    doc.add_heading('Codes de Réponse', 3)
    response_codes = [
        ('200', 'Succès'),
        ('204', 'Aucune donnée'),
        ('400', 'Paramètre manquant ou invalide'),
        ('403', 'Non authentifié'),
        ('409', 'Conflit (déjà confirmé)'),
        ('412', 'Precondition Failed (heartbeat absent)'),
        ('500', 'Erreur serveur')
    ]
    create_table_with_header(doc, ['Code', 'Signification'], response_codes)

    doc.add_page_break()

    # 2.5 Modèle de Données
    doc.add_heading('2.5 Modèle de Données', 2)

    doc.add_heading('Entités Principales', 3)
    entities = [
        ('TrainPass', 'Passage de train', 'train_pass_id, time_start, installation_id'),
        ('TrainPassCars', 'Wagons d\'un passage', 'car_type_id, train_pass_order, rfid_tag'),
        ('CarTypes', 'Types de wagons', 'icon_head, icon_tail, bogie_type, num_bogies'),
        ('BogieType', 'Types de bogies', 'svgname, axles_per_bogie'),
        ('SpringLocation', 'Position ressort', 'train_axle_number, cam_pos, confidence'),
        ('HumanConfirmations', 'Confirmations', 'present_not_absent, confirmed_by'),
        ('SpringImageLocation', 'Images ressort', 'image_path'),
        ('Installation', 'Sites d\'installation', 'location, description'),
        ('Users', 'Utilisateurs', 'user_name, password_hash, password_salt'),
        ('ConfidenceLevels', 'Niveaux confiance', 'conf_range, level_name')
    ]
    create_table_with_header(doc, ['Entité', 'Description', 'Champs Clés'], entities)

    doc.add_paragraph()
    doc.add_heading('Fonctions PostgreSQL', 3)
    pg_functions = [
        ('get_train_passes(timestamp, int, int?)', 'Récupère N passages avant une date'),
        ('get_pass_human_confirms_by_car_and_train_pass(int, int)', 'Wagons avec confirmations'),
        ('mark_human_confirmed(int, bool, str)', 'Enregistre une confirmation'),
        ('get_first_tag(int)', 'Premier tag RFID d\'un passage')
    ]
    create_table_with_header(doc, ['Fonction', 'Description'], pg_functions)

    doc.add_page_break()

    # =========================================================================
    # 3. AUDIT DE CODE
    # =========================================================================
    doc.add_heading('3. Audit de Code', 1)

    # 3.1 Analyse Statique
    doc.add_heading('3.1 Analyse Statique', 2)

    doc.add_heading('Résumé', 3)
    audit_summary = [
        ('Erreurs de syntaxe', '0', '✓ OK'),
        ('Problèmes de style (flake8)', '519', '⚠ À corriger'),
        ('Imports inutilisés', '7', '⚠ Minor'),
        ('Problèmes de sécurité', '3', '✗ Critique')
    ]
    create_table_with_header(doc, ['Métrique', 'Valeur', 'Statut'], audit_summary)

    doc.add_paragraph()
    doc.add_heading('Distribution des Erreurs de Style', 3)
    style_errors = [
        ('E231', 'Missing whitespace after \',\'', '118'),
        ('W293', 'Blank line contains whitespace', '60'),
        ('W291', 'Trailing whitespace', '59'),
        ('E265', 'Block comment format', '54'),
        ('E302', 'Expected 2 blank lines', '28'),
        ('E251', 'Spaces around \'=\'', '26'),
        ('E111', 'Indentation not multiple of 4', '23'),
        ('Autres', 'Divers problèmes mineurs', '151')
    ]
    create_table_with_header(doc, ['Code', 'Description', 'Occurrences'], style_errors)

    # 3.2 Audit Sécurité
    doc.add_heading('3.2 Audit de Sécurité', 2)

    doc.add_heading('Vulnérabilités Identifiées', 3)

    # Critical 1
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('🔴 CRITIQUE: Secret Key Hardcodée').bold = True
    doc.add_paragraph('Fichier: app.py ligne 27')
    doc.add_paragraph('Code actuel:')
    code1 = doc.add_paragraph("app.secret_key = 'slkfjaslkfdjlsadflknisdf64s6d4f56asf'")
    code1.runs[0].font.name = 'Courier New'
    doc.add_paragraph('Risque: Compromission des sessions si le code source est exposé.')
    doc.add_paragraph('Correction: Utiliser une variable d\'environnement.')

    # Critical 2
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('🔴 CRITIQUE: Credentials par Défaut').bold = True
    doc.add_paragraph('Fichier: user_management.py ligne 36')
    doc.add_paragraph('Code actuel:')
    code2 = doc.add_paragraph('create_user("eurotunnel", "Spr1ngs", "Euro Tunnel")')
    code2.runs[0].font.name = 'Courier New'
    doc.add_paragraph('Risque: Accès non autorisé avec des identifiants connus publiquement.')
    doc.add_paragraph('Correction: Ne pas créer d\'utilisateur par défaut en production.')

    # Medium
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('⚠️ MOYEN: Bare Except').bold = True
    doc.add_paragraph('Fichier: version.py ligne 8')
    doc.add_paragraph('Risque: Masque des erreurs inattendues.')
    doc.add_paragraph('Correction: Attraper des exceptions spécifiques.')

    doc.add_paragraph()
    doc.add_heading('Points Positifs', 3)
    positive_security = [
        'Hashage bcrypt des mots de passe avec salt unique',
        'Utilisation de SQLAlchemy ORM (protection SQL injection)',
        'Vérification de session sur chaque endpoint protégé',
        'Headers Cache-Control no-store sur les endpoints API',
        'Pas de stockage de mot de passe en clair'
    ]
    for item in positive_security:
        doc.add_paragraph('✓ ' + item, style='List Bullet')

    doc.add_page_break()

    # 3.3 Qualité du Code
    doc.add_heading('3.3 Qualité du Code', 2)

    doc.add_heading('Score par Module', 3)
    quality_scores = [
        ('redis_web.py', '⭐⭐⭐⭐', '⭐⭐⭐⭐', '⭐⭐⭐⭐', '88%'),
        ('missing_spring_endpoints.py', '⭐⭐⭐⭐', '⭐⭐⭐⭐', '⭐⭐⭐', '83%'),
        ('confidence_levels.py', '⭐⭐⭐', '⭐⭐⭐', '⭐⭐⭐⭐', '79%'),
        ('train_pass_endpoints.py', '⭐⭐⭐', '⭐⭐⭐', '⭐⭐⭐', '75%'),
        ('system_endpoints.py', '⭐⭐⭐', '⭐⭐⭐', '⭐⭐⭐', '75%'),
        ('app.py', '⭐⭐⭐', '⭐⭐⭐', '⭐⭐', '67%'),
        ('db_iface.py', '⭐⭐', '⭐⭐⭐', '⭐⭐⭐', '67%'),
        ('user_management.py', '⭐⭐⭐', '⭐⭐⭐', '⭐⭐', '67%')
    ]
    create_table_with_header(doc, ['Module', 'Lisibilité', 'Maintenabilité', 'Sécurité', 'Score'], quality_scores)

    doc.add_paragraph()
    doc.add_heading('Points Forts', 3)
    strengths = [
        'Architecture modulaire bien structurée',
        'Séparation claire des responsabilités (endpoints, db, services)',
        'Utilisation de type hints Python',
        'Validation des données avec Pydantic',
        'Gestion des sessions avec context manager',
        'Cache Redis pour optimiser les performances'
    ]
    for item in strengths:
        doc.add_paragraph('✓ ' + item, style='List Bullet')

    doc.add_heading('Points à Améliorer', 3)
    weaknesses = [
        'Formatage de code incohérent (519 violations flake8)',
        'Comparaisons à None avec == au lieu de is',
        'Imports inutilisés dans plusieurs fichiers',
        'Commentaires mal formatés',
        'Indentation parfois incorrecte',
        'Pas de tests unitaires (seulement intégration)'
    ]
    for item in weaknesses:
        doc.add_paragraph('✗ ' + item, style='List Bullet')

    doc.add_page_break()

    # =========================================================================
    # 4. RECOMMANDATIONS
    # =========================================================================
    doc.add_heading('4. Recommandations', 1)

    # 4.1 Matrice Impact vs Effort
    doc.add_heading('4.1 Matrice Impact vs Effort', 2)

    doc.add_paragraph(
        'La matrice suivante classe les recommandations selon leur impact sur la qualité/sécurité '
        'du système et l\'effort requis pour leur mise en œuvre.'
    )

    doc.add_paragraph()

    # Quick Wins
    doc.add_heading('🎯 Quick Wins (Impact Élevé, Effort Faible)', 3)
    quick_wins = [
        ('SEC-01', 'Externaliser la secret key', 'CRITIQUE', '30 min', 'Variable d\'environnement'),
        ('SEC-02', 'Supprimer credentials par défaut', 'CRITIQUE', '15 min', 'Supprimer create_users_if_none en prod'),
        ('QUA-01', 'Formater avec Black', 'MOYEN', '5 min', 'black eurotunnel_web/'),
        ('QUA-02', 'Trier imports avec isort', 'FAIBLE', '5 min', 'isort eurotunnel_web/')
    ]
    create_table_with_header(doc, ['ID', 'Action', 'Impact', 'Effort', 'Détail'], quick_wins)

    doc.add_paragraph()

    # Major Projects
    doc.add_heading('📊 Projets Majeurs (Impact Élevé, Effort Élevé)', 3)
    major_projects = [
        ('SEC-03', 'Ajouter rate limiting', 'ÉLEVÉ', '2-4h', 'Flask-Limiter sur /login'),
        ('SEC-04', 'Implémenter CSRF protection', 'ÉLEVÉ', '4-8h', 'Flask-WTF'),
        ('QUA-03', 'Ajouter tests unitaires', 'ÉLEVÉ', '2-3j', 'pytest, mocks'),
        ('ARC-01', 'Migrer vers Python 3.12', 'MOYEN', '1-2h', 'Compatibilité pyproject.toml')
    ]
    create_table_with_header(doc, ['ID', 'Action', 'Impact', 'Effort', 'Détail'], major_projects)

    doc.add_paragraph()

    # Fill-ins
    doc.add_heading('📝 Améliorations Continues (Impact Moyen, Effort Faible)', 3)
    fill_ins = [
        ('QUA-04', 'Corriger comparaisons None', 'FAIBLE', '30 min', 'is None au lieu de == None'),
        ('QUA-05', 'Supprimer imports inutilisés', 'FAIBLE', '15 min', 'autoflake'),
        ('QUA-06', 'Corriger bare except', 'FAIBLE', '15 min', 'version.py'),
        ('DOC-01', 'Ajouter docstrings manquants', 'FAIBLE', '2-4h', 'Fonctions publiques')
    ]
    create_table_with_header(doc, ['ID', 'Action', 'Impact', 'Effort', 'Détail'], fill_ins)

    doc.add_paragraph()

    # Maybe Later
    doc.add_heading('📅 À Considérer (Impact Faible, Effort Élevé)', 3)
    maybe_later = [
        ('ARC-02', 'Migrer Vue.js 2 vers 3', 'FAIBLE', '1-2 sem', 'Breaking changes'),
        ('ARC-03', 'Ajouter API versioning', 'FAIBLE', '1-2j', '/api/v1/'),
        ('QUA-07', 'Ajouter type checking mypy', 'FAIBLE', '1-2j', 'Configuration stricte')
    ]
    create_table_with_header(doc, ['ID', 'Action', 'Impact', 'Effort', 'Détail'], maybe_later)

    doc.add_page_break()

    # 4.2 Plan d'Action
    doc.add_heading('4.2 Plan d\'Action Priorisé', 2)

    doc.add_heading('Phase 1: Sécurité Immédiate (Semaine 1)', 3)
    phase1 = [
        ('1.1', 'Externaliser app.secret_key vers variable d\'environnement', 'SEC-01', '30 min'),
        ('1.2', 'Désactiver création user par défaut en production', 'SEC-02', '15 min'),
        ('1.3', 'Documenter les credentials nécessaires', '-', '30 min'),
        ('1.4', 'Valider le déploiement', '-', '1h')
    ]
    create_table_with_header(doc, ['#', 'Tâche', 'Réf', 'Durée'], phase1)

    doc.add_paragraph()
    doc.add_heading('Phase 2: Qualité de Code (Semaine 2)', 3)
    phase2 = [
        ('2.1', 'Installer et configurer Black, isort, flake8', 'QUA-01/02', '30 min'),
        ('2.2', 'Formater l\'ensemble du code', 'QUA-01/02', '15 min'),
        ('2.3', 'Corriger les 7 comparaisons None', 'QUA-04', '30 min'),
        ('2.4', 'Supprimer les imports inutilisés', 'QUA-05', '15 min'),
        ('2.5', 'Corriger le bare except', 'QUA-06', '15 min'),
        ('2.6', 'Ajouter pre-commit hooks', '-', '30 min')
    ]
    create_table_with_header(doc, ['#', 'Tâche', 'Réf', 'Durée'], phase2)

    doc.add_paragraph()
    doc.add_heading('Phase 3: Sécurité Renforcée (Semaines 3-4)', 3)
    phase3 = [
        ('3.1', 'Implémenter rate limiting sur /login', 'SEC-03', '2-4h'),
        ('3.2', 'Ajouter protection CSRF', 'SEC-04', '4-8h'),
        ('3.3', 'Configurer HTTPS obligatoire', '-', '2h'),
        ('3.4', 'Ajouter audit logging des confirmations', '-', '4h'),
        ('3.5', 'Tests de pénétration basiques', '-', '4h')
    ]
    create_table_with_header(doc, ['#', 'Tâche', 'Réf', 'Durée'], phase3)

    doc.add_paragraph()
    doc.add_heading('Phase 4: Tests et Documentation (Mois 2)', 3)
    phase4 = [
        ('4.1', 'Écrire tests unitaires pour db_iface', 'QUA-03', '1j'),
        ('4.2', 'Écrire tests unitaires pour endpoints', 'QUA-03', '1j'),
        ('4.3', 'Configurer couverture de code (coverage)', '-', '2h'),
        ('4.4', 'Ajouter docstrings aux fonctions publiques', 'DOC-01', '4h'),
        ('4.5', 'Mettre à jour README avec instructions', '-', '2h')
    ]
    create_table_with_header(doc, ['#', 'Tâche', 'Réf', 'Durée'], phase4)

    doc.add_paragraph()
    doc.add_paragraph()

    # Conclusion
    doc.add_heading('Conclusion', 2)
    doc.add_paragraph(
        'L\'application VICORE présente une architecture solide et des fonctionnalités bien implémentées. '
        'Les principales préoccupations concernent la sécurité (secret key, credentials par défaut) qui '
        'doivent être adressées immédiatement avant tout déploiement en production. '
        'Les problèmes de qualité de code sont nombreux mais essentiellement cosmétiques et peuvent être '
        'corrigés automatiquement avec des outils de formatage.'
    )

    doc.add_paragraph()
    p = doc.add_paragraph('Effort total estimé pour les corrections prioritaires: ')
    p.add_run('2-3 semaines').bold = True
    p.add_run(' (temps développeur)')

    # Save
    output_path = 'C:/Users/lengl/Documents/eurotunnel_web-master/VICORE_Rapport_Complet.docx'
    doc.save(output_path)
    print(f"Rapport généré: {output_path}")
    return output_path

if __name__ == '__main__':
    generate_report()
